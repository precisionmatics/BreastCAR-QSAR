#!/usr/bin/env python3
"""
Generate manuscript Figures 1–10 and embed them in the revised docx.
NO embedded figure titles — captions live in the manuscript text only.
"""

import os
import glob
import json
import warnings
import numpy as np
import pandas as pd
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.ticker as ticker
from matplotlib.patches import Patch
from scipy.stats import pearsonr
from docx import Document

warnings.filterwarnings('ignore')

# ── Paths ──────────────────────────────────────────────────────────────────────
BASE      = '/home/stalin/Desktop/Breast_Cancer_ML'
NEW_DIR   = os.path.join(BASE, 'NEW')
RES_DIR   = os.path.join(BASE, 'revised_results', 'results')
REV_DIR   = os.path.join(BASE, 'revised_results')
FIG_DIR   = os.path.join(REV_DIR, 'figures')
DOCX_PATH = os.path.join(REV_DIR, '20_May_2026_Manuscript_Revised.docx')
BEST_CSV  = os.path.join(REV_DIR, 'best_model_summary.csv')
REG_CSV   = os.path.join(REV_DIR, 'regression_summary.csv')

os.makedirs(FIG_DIR, exist_ok=True)

ALG_COLORS = {'RF':'#4C72B0','XGB':'#DD8452','LGB':'#55A868','SVM':'#C44E52','MLP':'#8172B2'}
ALGS = ['RF', 'XGB', 'LGB', 'SVM', 'MLP']

SKIP_DIRS = {'ALL_TXT', 'Non_BC_Targets', 'Others'}

def pub_style(ax):
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)

def panel_label(ax, label, x=-0.08, y=1.04):
    ax.text(x, y, label, transform=ax.transAxes,
            fontsize=13, fontweight='bold', va='top', ha='right')

# ── Load data ──────────────────────────────────────────────────────────────────
bms = pd.read_csv(BEST_CSV).set_index('Target')
reg_full = pd.read_csv(REG_CSV)

def load_all_csv_pIC50():
    target_data = {}
    for subdir in sorted(os.listdir(NEW_DIR)):
        if subdir in SKIP_DIRS:
            continue
        dpath = os.path.join(NEW_DIR, subdir)
        if not os.path.isdir(dpath):
            continue
        csvs = [f for f in glob.glob(os.path.join(dpath, '*.csv'))
                if not os.path.basename(f).startswith('._')]
        if not csvs:
            continue
        frames = []
        for csv in csvs:
            try:
                df = pd.read_csv(csv, low_memory=False)
            except Exception:
                continue
            pic50 = None
            for col in df.columns:
                if col.strip() == 'pIC50':
                    pic50 = df[col].dropna().astype(float); break
                if col.strip() == 'pIC50_value':
                    pic50 = df[col].dropna().astype(float); break
            if pic50 is None:
                for col in df.columns:
                    if 'Standard Value' in col:
                        vals = pd.to_numeric(df[col], errors='coerce').dropna()
                        vals = vals[vals > 0]
                        if len(vals):
                            pic50 = -np.log10(vals * 1e-9)
                        break
            if pic50 is not None and len(pic50) > 0:
                frames.append(pic50)
        if frames:
            target_data[subdir] = pd.concat(frames, ignore_index=True)
    return target_data

print("Loading pIC50 data from NEW/ ...")
target_data = load_all_csv_pIC50()
all_pic50 = pd.concat(list(target_data.values()), ignore_index=True)
print(f"  {len(all_pic50):,} values across {len(target_data)} targets")

# ═══════════════════════════════════════════════════════════════════════════════
# FIGURE 1 — Global pIC50 distribution
# ═══════════════════════════════════════════════════════════════════════════════
print("\n── Figure 1: Global pIC50 distribution ──")
fig1, ax1 = plt.subplots(figsize=(10, 5))
ax1.hist(all_pic50, bins=60, color='steelblue', edgecolor='white', linewidth=0.4, alpha=0.85)
ax1.axvline(7.0, color='#2CA02C', linestyle='--', linewidth=1.5,
            label='pIC50 = 7.0 (Active threshold)')
ax1.axvline(5.0, color='#D62728', linestyle='--', linewidth=1.5,
            label='pIC50 = 5.0 (Inactive threshold)')
ax1.set_xlabel('pIC50', fontsize=12)
ax1.set_ylabel('Number of Compounds', fontsize=12)
ax1.legend(fontsize=10)
ax1.yaxis.set_major_formatter(ticker.FuncFormatter(lambda x, _: f'{int(x):,}'))
pub_style(ax1)
plt.tight_layout()
fig1_path = os.path.join(FIG_DIR, 'manuscript_fig1.png')
fig1.savefig(fig1_path, dpi=200, bbox_inches='tight')
plt.close(fig1)
print(f"  Saved: {fig1_path}")

# ═══════════════════════════════════════════════════════════════════════════════
# FIGURE 2 — Dataset overview (4 panels)
# ═══════════════════════════════════════════════════════════════════════════════
print("\n── Figure 2: Dataset overview ──")

sorted_by_n   = bms['N'].sort_values(ascending=False)
sorted_targets = sorted_by_n.index.tolist()

counts = {}
for t, s in target_data.items():
    counts[t] = {
        'Active':   int((s >= 7.0).sum()),
        'Moderate': int(((s >= 5.0) & (s < 7.0)).sum()),
        'Inactive': int((s < 5.0).sum()),
    }
counts_df = pd.DataFrame(counts).T
common = [t for t in sorted_targets if t in counts_df.index]
x = np.arange(len(common))

fig2, axes2 = plt.subplots(2, 2, figsize=(18, 14))
ax2A, ax2B, ax2C, ax2D = axes2[0,0], axes2[0,1], axes2[1,0], axes2[1,1]

# A — compound count
n_vals = [bms.loc[t, 'N'] if t in bms.index else 0 for t in common]
ax2A.bar(x, n_vals, color='#4C72B0', width=0.7)
ax2A.set_xticks(x); ax2A.set_xticklabels(common, rotation=90, fontsize=7)
ax2A.set_ylabel('Number of Compounds', fontsize=10)
ax2A.yaxis.set_major_formatter(ticker.FuncFormatter(lambda v, _: f'{int(v):,}'))
panel_label(ax2A, 'A')
pub_style(ax2A)

# B — stacked activity classes
a_v = [counts_df.loc[t, 'Active']   if t in counts_df.index else 0 for t in common]
m_v = [counts_df.loc[t, 'Moderate'] if t in counts_df.index else 0 for t in common]
i_v = [counts_df.loc[t, 'Inactive'] if t in counts_df.index else 0 for t in common]
ax2B.bar(x, a_v, color='#4C72B0', width=0.7, label='Active (pIC50 ≥ 7)')
ax2B.bar(x, m_v, bottom=a_v, color='#DD8452', width=0.7, label='Moderate (5 ≤ pIC50 < 7)')
ax2B.bar(x, i_v, bottom=[a+m for a,m in zip(a_v,m_v)], color='#55A868', width=0.7,
         label='Inactive (pIC50 < 5)')
ax2B.set_xticks(x); ax2B.set_xticklabels(common, rotation=90, fontsize=7)
ax2B.set_ylabel('Number of Compounds', fontsize=10)
ax2B.legend(fontsize=8, loc='upper right')
panel_label(ax2B, 'B')
pub_style(ax2B)

# C — pIC50 mean±SD (top 20)
top20 = [t for t in sorted_targets if t in target_data][:20]
means = [target_data[t].mean() for t in top20]
stds  = [target_data[t].std()  for t in top20]
x20   = np.arange(len(top20))
ax2C.bar(x20, means, yerr=stds, capsize=3, color='#5975A4', width=0.65,
         error_kw={'elinewidth':1, 'ecolor':'gray'})
ax2C.axhline(7.0, color='#2CA02C', linestyle='--', linewidth=1.2, label='pIC50 = 7.0')
ax2C.axhline(5.0, color='#D62728', linestyle='--', linewidth=1.2, label='pIC50 = 5.0')
ax2C.set_xticks(x20); ax2C.set_xticklabels(top20, rotation=45, fontsize=8, ha='right')
ax2C.set_ylabel('pIC50 (mean ± SD)', fontsize=10)
ax2C.legend(fontsize=8)
panel_label(ax2C, 'C')
pub_style(ax2C)

# D — AD coverage
ad_targets = [t for t in common if 'AD_pct_in' in bms.columns and t in bms.index]
ad_vals    = [bms.loc[t, 'AD_pct_in'] for t in ad_targets]
x_ad = np.arange(len(ad_targets))
ad_colors = ['#C44E52' if v < 80 else '#4C72B0' for v in ad_vals]
ax2D.bar(x_ad, ad_vals, color=ad_colors, width=0.7)
ax2D.axhline(80, color='orange', linestyle='--', linewidth=1.2, label='80% threshold')
ax2D.set_xticks(x_ad); ax2D.set_xticklabels(ad_targets, rotation=90, fontsize=7)
ax2D.set_ylabel('% Test Set in AD', fontsize=10)
ax2D.set_ylim(0, 105)
ax2D.legend(fontsize=8)
panel_label(ax2D, 'D')
pub_style(ax2D)

plt.tight_layout()
fig2_path = os.path.join(FIG_DIR, 'manuscript_fig2.png')
fig2.savefig(fig2_path, dpi=200, bbox_inches='tight')
plt.close(fig2)
print(f"  Saved: {fig2_path}")

# ═══════════════════════════════════════════════════════════════════════════════
# FIGURE 3 — QSAR performance overview (4 panels)
# ═══════════════════════════════════════════════════════════════════════════════
print("\n── Figure 3: QSAR performance overview ──")

bms_r = bms.copy().reset_index()
bms_r2_sorted = bms_r.sort_values('R2', ascending=False)
targets_r2 = bms_r2_sorted['Target'].tolist()
r2_vals    = bms_r2_sorted['R2'].tolist()
rmse_vals  = bms_r2_sorted['RMSE'].tolist()
f1_vals    = bms_r2_sorted['F1_weighted'].tolist()
x3 = np.arange(len(targets_r2))

reg_freq = {a: 0 for a in ALGS}
clf_freq = {a: 0 for a in ALGS}
for _, row in bms_r.iterrows():
    if row.get('Best_Reg_Model') in reg_freq: reg_freq[row['Best_Reg_Model']] += 1
    if row.get('Best_Clf_Model') in clf_freq: clf_freq[row['Best_Clf_Model']] += 1

fig3, axes3 = plt.subplots(2, 2, figsize=(18, 14))
ax3A, ax3B, ax3C, ax3D = axes3[0,0], axes3[0,1], axes3[1,0], axes3[1,1]

legend_els = [Patch(facecolor=ALG_COLORS[a], label=a) for a in ALGS]

# A — R²
colors_a = [ALG_COLORS.get(m, 'gray') for m in bms_r2_sorted['Best_Reg_Model']]
ax3A.bar(x3, r2_vals, color=colors_a, width=0.7)
ax3A.set_xticks(x3); ax3A.set_xticklabels(targets_r2, rotation=90, fontsize=7)
ax3A.set_ylabel('R²', fontsize=10); ax3A.set_ylim(0, 1.0)
ax3A.legend(handles=legend_els, fontsize=8, loc='upper right')
panel_label(ax3A, 'A'); pub_style(ax3A)

# B — RMSE
ax3B.bar(x3, rmse_vals, color=colors_a, width=0.7)
ax3B.set_xticks(x3); ax3B.set_xticklabels(targets_r2, rotation=90, fontsize=7)
ax3B.set_ylabel('RMSE', fontsize=10)
ax3B.legend(handles=legend_els, fontsize=8, loc='upper right')
panel_label(ax3B, 'B'); pub_style(ax3B)

# C — F1 weighted
colors_c = [ALG_COLORS.get(m, 'gray') for m in bms_r2_sorted['Best_Clf_Model']]
ax3C.bar(x3, f1_vals, color=colors_c, width=0.7)
ax3C.set_xticks(x3); ax3C.set_xticklabels(targets_r2, rotation=90, fontsize=7)
ax3C.set_ylabel('F1 (weighted)', fontsize=10); ax3C.set_ylim(0, 1.0)
clf_legend = [Patch(facecolor=ALG_COLORS[a], label=a) for a in ALGS]
ax3C.legend(handles=clf_legend, fontsize=8, loc='upper right')
panel_label(ax3C, 'C'); pub_style(ax3C)

# D — Algorithm frequency
x_alg = np.arange(len(ALGS)); w = 0.35
reg_c = [reg_freq[a] for a in ALGS]
clf_c = [clf_freq[a] for a in ALGS]
b1 = ax3D.bar(x_alg - w/2, reg_c, width=w,
              color=[ALG_COLORS[a] for a in ALGS], label='Regression')
b2 = ax3D.bar(x_alg + w/2, clf_c, width=w,
              color=[ALG_COLORS[a] for a in ALGS], alpha=0.55,
              edgecolor='black', linewidth=0.8, label='Classification')
ax3D.set_xticks(x_alg); ax3D.set_xticklabels(ALGS, fontsize=10)
ax3D.set_ylabel('Number of Targets', fontsize=10)
for bar in list(b1) + list(b2):
    h = bar.get_height()
    if h > 0:
        ax3D.text(bar.get_x() + bar.get_width()/2, h + 0.15, str(int(h)),
                  ha='center', va='bottom', fontsize=8)
ax3D.legend(fontsize=9)
panel_label(ax3D, 'D'); pub_style(ax3D)

plt.tight_layout()
fig3_path = os.path.join(FIG_DIR, 'manuscript_fig3.png')
fig3.savefig(fig3_path, dpi=200, bbox_inches='tight')
plt.close(fig3)
print(f"  Saved: {fig3_path}")

# ═══════════════════════════════════════════════════════════════════════════════
# FIGURE 4 — Performance heatmap (4 metrics × 53 targets)
# ═══════════════════════════════════════════════════════════════════════════════
print("\n── Figure 4: Performance heatmap ──")

metrics4 = ['R2', 'F1_weighted', 'MCC', 'Balanced_Accuracy']
metric_labels = ['R²', 'F1 (weighted)', 'MCC', 'Balanced Accuracy']
bms_hm = bms.sort_values('R2', ascending=False)
hm_targets = bms_hm.index.tolist()
hm_data = bms_hm[metrics4].T.values.astype(float)

fig4, ax4 = plt.subplots(figsize=(22, 4.5))
im = ax4.imshow(hm_data, cmap='RdYlGn', vmin=0, vmax=1, aspect='auto')

ax4.set_xticks(np.arange(len(hm_targets)))
ax4.set_xticklabels(hm_targets, rotation=90, fontsize=7)
ax4.set_yticks(np.arange(len(metrics4)))
ax4.set_yticklabels(metric_labels, fontsize=10)

for i in range(len(metrics4)):
    for j in range(len(hm_targets)):
        val = hm_data[i, j]
        txt = f'{val:.2f}' if not np.isnan(val) else 'N/A'
        nv  = val if not np.isnan(val) else 0.5
        bg  = plt.cm.RdYlGn(nv)[:3]
        lum = 0.299*bg[0] + 0.587*bg[1] + 0.114*bg[2]
        ax4.text(j, i, txt, ha='center', va='center',
                 fontsize=4.5, color='black' if lum > 0.45 else 'white')

plt.colorbar(im, ax=ax4, fraction=0.015, pad=0.01, label='Score')
plt.tight_layout()
fig4_path = os.path.join(FIG_DIR, 'manuscript_fig4.png')
fig4.savefig(fig4_path, dpi=200, bbox_inches='tight')
plt.close(fig4)
print(f"  Saved: {fig4_path}")

# ═══════════════════════════════════════════════════════════════════════════════
# FIGURE 5 — Pairwise R² scatter (RF vs XGB, RF vs LGB, XGB vs LGB)
# ═══════════════════════════════════════════════════════════════════════════════
print("\n── Figure 5: Pairwise R² scatter ──")

pivot5 = reg_full.pivot_table(index='Target', columns='Model', values='R2')
pairs5 = [('RF','XGB'), ('RF','LGB'), ('XGB','LGB')]
panel_labels5 = ['A', 'B', 'C']

fig5, axes5 = plt.subplots(1, 3, figsize=(15, 5))

for ax, (m1, m2), pl in zip(axes5, pairs5, panel_labels5):
    sub = pivot5[[m1, m2]].dropna()
    xv, yv = sub[m1].values, sub[m2].values
    labels = sub.index.tolist()
    r, pval = pearsonr(xv, yv)

    ax.scatter(xv, yv, color='steelblue', s=40, alpha=0.8, zorder=3)
    for lbl, xi, yi in zip(labels, xv, yv):
        ax.annotate(lbl, (xi, yi), fontsize=5, alpha=0.7,
                    xytext=(2, 2), textcoords='offset points')

    mn = min(xv.min(), yv.min()) - 0.05
    mx = max(xv.max(), yv.max()) + 0.05
    ax.plot([mn, mx], [mn, mx], 'r--', linewidth=1.2, zorder=2)
    ax.set_xlim(mn, mx); ax.set_ylim(mn, mx)
    ax.set_xlabel(f'{m1}  R²', fontsize=10)
    ax.set_ylabel(f'{m2}  R²', fontsize=10)
    ax.text(0.05, 0.95, f'r = {r:.3f}', transform=ax.transAxes,
            fontsize=9, va='top', ha='left',
            bbox=dict(boxstyle='round,pad=0.3', fc='white', ec='gray', alpha=0.8))
    ax.set_aspect('equal'); ax.grid(True, alpha=0.3)
    pub_style(ax)
    panel_label(ax, pl)

plt.tight_layout()
fig5_path = os.path.join(FIG_DIR, 'manuscript_fig5.png')
fig5.savefig(fig5_path, dpi=200, bbox_inches='tight')
plt.close(fig5)
print(f"  Saved: {fig5_path}")

# ═══════════════════════════════════════════════════════════════════════════════
# FIGURE 6 — BCL2 all-model regression metrics
# ═══════════════════════════════════════════════════════════════════════════════
print("\n── Figure 6: BCL2 regression model comparison ──")

bcl2_d = json.load(open(os.path.join(RES_DIR, 'BCL2_results.json')))
bcl2_reg = {m: bcl2_d['reg'][m] for m in ALGS if m in bcl2_d['reg']}
bcl2_clf = {m: bcl2_d['clf'][m] for m in ALGS if m in bcl2_d['clf']}
bcl2_meta = bcl2_d['meta']

reg_metrics = ['r2', 'rmse', 'mae', 'cv_r2_mean']
reg_ylabels = ['R²', 'RMSE', 'MAE', 'CV R² (mean)']
alg_names = [m for m in ALGS if m in bcl2_reg]
x6 = np.arange(len(alg_names))
w6 = 0.18

fig6, axes6 = plt.subplots(1, 4, figsize=(16, 5))

for ax, met, ylabel in zip(axes6, reg_metrics, reg_ylabels):
    vals = [bcl2_reg[m].get(met, 0) for m in alg_names]
    bars = ax.bar(x6, vals, color=[ALG_COLORS[m] for m in alg_names], width=0.55)
    ax.set_xticks(x6); ax.set_xticklabels(alg_names, fontsize=10)
    ax.set_ylabel(ylabel, fontsize=10)
    for bar, v in zip(bars, vals):
        ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.005,
                f'{v:.3f}', ha='center', va='bottom', fontsize=8)
    pub_style(ax)

axes6[0].set_ylim(0, 1.0)
axes6[3].set_ylim(0, 1.0)
panel_label(axes6[0], 'A'); panel_label(axes6[1], 'B')
panel_label(axes6[2], 'C'); panel_label(axes6[3], 'D')
plt.tight_layout()
fig6_path = os.path.join(FIG_DIR, 'manuscript_fig6.png')
fig6.savefig(fig6_path, dpi=200, bbox_inches='tight')
plt.close(fig6)
print(f"  Saved: {fig6_path}")

# ═══════════════════════════════════════════════════════════════════════════════
# FIGURE 7 — BCL2 classification: per-model metrics + per-class F1
# ═══════════════════════════════════════════════════════════════════════════════
print("\n── Figure 7: BCL2 classification model comparison ──")

clf_met_keys   = ['f1_weighted', 'mcc', 'balanced_accuracy', 'auc_roc']
clf_met_labels = ['F1 (weighted)', 'MCC', 'Balanced Accuracy', 'AUC-ROC']

fig7, axes7 = plt.subplots(1, 2, figsize=(16, 6))

# Panel A — 4 metrics × 5 models grouped bar
ax7A = axes7[0]
x7 = np.arange(len(alg_names))
w7 = 0.18
offsets = np.linspace(-(len(clf_met_keys)-1)/2, (len(clf_met_keys)-1)/2, len(clf_met_keys)) * w7
metric_colors = ['#4C72B0','#DD8452','#55A868','#C44E52']
for offset, mkey, mlab, mc in zip(offsets, clf_met_keys, clf_met_labels, metric_colors):
    vals = [bcl2_clf[m].get(mkey, 0) for m in alg_names]
    ax7A.bar(x7 + offset, vals, width=w7, label=mlab, color=mc, alpha=0.85)
ax7A.set_xticks(x7); ax7A.set_xticklabels(alg_names, fontsize=10)
ax7A.set_ylabel('Score', fontsize=10); ax7A.set_ylim(0, 1.05)
ax7A.axhline(0.8, color='gray', ls='--', lw=0.8, alpha=0.5)
ax7A.legend(fontsize=8, loc='lower right')
panel_label(ax7A, 'A'); pub_style(ax7A)

# Panel B — per-class F1
ax7B = axes7[1]
class_keys   = ['f1_Active', 'f1_Moderate', 'f1_Inactive']
class_labels = ['Active', 'Moderate', 'Inactive']
class_colors = ['#4C72B0', '#DD8452', '#55A868']
w7b = 0.22
offsets_b = np.array([-1, 0, 1]) * w7b
for offset, ckey, clab, cc in zip(offsets_b, class_keys, class_labels, class_colors):
    vals = [bcl2_clf[m].get(ckey, 0) for m in alg_names]
    ax7B.bar(x7 + offset, vals, width=w7b, label=clab, color=cc, alpha=0.85)
ax7B.set_xticks(x7); ax7B.set_xticklabels(alg_names, fontsize=10)
ax7B.set_ylabel('F1 Score', fontsize=10); ax7B.set_ylim(0, 1.05)
ax7B.legend(fontsize=9, loc='lower right')
panel_label(ax7B, 'B'); pub_style(ax7B)

plt.tight_layout()
fig7_path = os.path.join(FIG_DIR, 'manuscript_fig7.png')
fig7.savefig(fig7_path, dpi=200, bbox_inches='tight')
plt.close(fig7)
print(f"  Saved: {fig7_path}")

# ═══════════════════════════════════════════════════════════════════════════════
# FIGURE 8 — BCL2 validation: Y-randomization + scaffold split
# ═══════════════════════════════════════════════════════════════════════════════
print("\n── Figure 8: BCL2 validation ──")

best_reg_m = max(alg_names, key=lambda m: bcl2_reg[m].get('r2', -999))
best_clf_m = max(alg_names, key=lambda m: bcl2_clf[m].get('f1_weighted', -999))
real_r2    = bcl2_reg[best_reg_m]['r2']
real_f1    = bcl2_clf[best_clf_m]['f1_weighted']
yrand_r2   = bcl2_meta['yrandom_reg_mean']
yrand_f1   = bcl2_meta['yrandom_clf_mean']
scaff_r2   = bcl2_meta['scaffold_reg_r2']
scaff_f1   = bcl2_meta['scaffold_clf_f1']
ad_pct     = bcl2_meta['ad_pct_test_in_ad']

fig8, axes8 = plt.subplots(1, 3, figsize=(15, 5))

# Panel A — Y-randomization
ax8A = axes8[0]
cats   = ['Real Model', 'Y-Randomized']
r2s    = [real_r2, yrand_r2]
f1s    = [real_f1, yrand_f1]
xA = np.arange(2); wA = 0.3
b_r2 = ax8A.bar(xA - wA/2, r2s, wA, color=['#4C72B0','#C44E52'], label='R² (Regression)', alpha=0.85)
b_f1 = ax8A.bar(xA + wA/2, f1s, wA, color=['#4C72B0','#C44E52'], label='F1 (Classification)', alpha=0.55, edgecolor='black', lw=0.8)
ax8A.set_xticks(xA); ax8A.set_xticklabels(cats, fontsize=10)
ax8A.set_ylabel('Score', fontsize=10); ax8A.set_ylim(-0.2, 1.1)
ax8A.axhline(0, color='black', lw=0.8, ls='--')
for bar, v in zip(list(b_r2)+list(b_f1), r2s+f1s):
    ax8A.text(bar.get_x()+bar.get_width()/2, bar.get_height()+0.02,
              f'{v:.3f}', ha='center', va='bottom', fontsize=9)
ax8A.legend(fontsize=8)
panel_label(ax8A, 'A'); pub_style(ax8A)

# Panel B — Scaffold vs Random split
ax8B = axes8[1]
cats_sc = ['Random Split', 'Scaffold Split']
r2_sc   = [real_r2, scaff_r2]
f1_sc   = [real_f1, scaff_f1]
xB = np.arange(2)
bB_r = ax8B.bar(xB - wA/2, r2_sc, wA, color=['#4C72B0','#55A868'], label='R² (Regression)', alpha=0.85)
bB_f = ax8B.bar(xB + wA/2, f1_sc, wA, color=['#4C72B0','#55A868'], label='F1 (Classification)', alpha=0.55, edgecolor='black', lw=0.8)
ax8B.set_xticks(xB); ax8B.set_xticklabels(cats_sc, fontsize=10)
ax8B.set_ylabel('Score', fontsize=10); ax8B.set_ylim(0, 1.1)
for bar, v in zip(list(bB_r)+list(bB_f), r2_sc+f1_sc):
    ax8B.text(bar.get_x()+bar.get_width()/2, bar.get_height()+0.01,
              f'{v:.3f}', ha='center', va='bottom', fontsize=9)
ax8B.legend(fontsize=8)
panel_label(ax8B, 'B'); pub_style(ax8B)

# Panel C — AD coverage gauge
ax8C = axes8[2]
ax8C.barh(['Test Set\nCoverage'], [ad_pct], color='#4C72B0', height=0.4)
ax8C.barh(['Test Set\nCoverage'], [100-ad_pct], left=[ad_pct],
           color='#C44E52', height=0.4, alpha=0.5, label='Outside AD')
ax8C.axvline(80, color='orange', ls='--', lw=1.5, label='80% threshold')
ax8C.set_xlim(0, 105); ax8C.set_xlabel('% Compounds', fontsize=10)
ax8C.text(ad_pct/2, 0, f'{ad_pct:.1f}%', ha='center', va='center',
          color='white', fontweight='bold', fontsize=12)
ax8C.legend(fontsize=8, loc='lower right')
panel_label(ax8C, 'C'); pub_style(ax8C)

plt.tight_layout()
fig8_path = os.path.join(FIG_DIR, 'manuscript_fig8.png')
fig8.savefig(fig8_path, dpi=200, bbox_inches='tight')
plt.close(fig8)
print(f"  Saved: {fig8_path}")

# ═══════════════════════════════════════════════════════════════════════════════
# FIGURES 9 & 10 — Web app screenshots (Selenium + headless Firefox)
# ═══════════════════════════════════════════════════════════════════════════════
print("\n── Figures 9 & 10: Web app screenshots ──")

fig9_path  = os.path.join(FIG_DIR, 'manuscript_fig9.png')
fig10_path = os.path.join(FIG_DIR, 'manuscript_fig10.png')

try:
    import subprocess, time, shutil
    from selenium import webdriver
    from selenium.webdriver.firefox.options import Options
    from selenium.webdriver.firefox.service import Service

    # Check geckodriver
    gecko = shutil.which('geckodriver') or '/snap/bin/geckodriver'
    if not os.path.exists(gecko):
        raise FileNotFoundError(f"geckodriver not found at {gecko}")

    # Start uvicorn
    webapp_dir = os.path.join(BASE, 'webapp')
    proc = subprocess.Popen(
        ['python3', '-m', 'uvicorn', 'app:app', '--host', '127.0.0.1', '--port', '8765'],
        cwd=webapp_dir,
        stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL
    )
    time.sleep(4)

    opts = Options()
    opts.add_argument('--headless')
    opts.add_argument('--width=1280')
    opts.add_argument('--height=900')
    svc = Service(executable_path=gecko)
    driver = webdriver.Firefox(service=svc, options=opts)
    driver.set_window_size(1280, 900)

    # Fig 9 — main prediction page
    driver.get('http://127.0.0.1:8765/')
    time.sleep(3)
    driver.save_screenshot(fig9_path)
    print(f"  Saved: {fig9_path}")

    # Fig 10 — scroll or navigate to results / second view
    driver.execute_script("window.scrollTo(0, 400)")
    time.sleep(1.5)
    driver.save_screenshot(fig10_path)
    print(f"  Saved: {fig10_path}")

    driver.quit()
    proc.terminate()
    print("  Web app screenshots captured.")

except Exception as e:
    print(f"  [WARN] Web app screenshot failed: {e}")
    print("  Creating placeholder figures for 9 & 10...")

    # Placeholder with informative text
    for path, title in [(fig9_path, 'Web Application — Prediction Interface'),
                        (fig10_path, 'Web Application — Results View')]:
        fig_p, ax_p = plt.subplots(figsize=(12, 7))
        ax_p.text(0.5, 0.5, f'[Screenshot: {title}]\n\nStart the web app with:\nuvicorn app:app --host 0.0.0.0 --port 8000\nthen run generate_manuscript_figures.py again',
                  ha='center', va='center', fontsize=14,
                  transform=ax_p.transAxes,
                  bbox=dict(boxstyle='round', fc='lightyellow', ec='orange'))
        ax_p.axis('off')
        fig_p.savefig(path, dpi=150, bbox_inches='tight')
        plt.close(fig_p)
        print(f"  Placeholder saved: {path}")

# ═══════════════════════════════════════════════════════════════════════════════
# EMBED ALL 10 FIGURES INTO DOCX
# ═══════════════════════════════════════════════════════════════════════════════
print("\n── Embedding figures in docx ──")

def replace_image_in_docx(docx_path, para_idx, new_image_path):
    doc = Document(docx_path)
    para = doc.paragraphs[para_idx]
    drawings = para._element.findall(
        './/{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing')
    if not drawings:
        print(f"  [WARN] No drawing at para {para_idx}")
        return False
    blip = drawings[0].find(
        './/{http://schemas.openxmlformats.org/drawingml/2006/main}blip')
    if blip is None:
        print(f"  [WARN] No blip at para {para_idx}")
        return False
    r_embed = blip.get(
        '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
    image_part = doc.part.rels[r_embed].target_part
    with open(new_image_path, 'rb') as f:
        image_part._blob = f.read()
    doc.save(docx_path)
    return True

replacements = [
    (76,  fig1_path,  'Figure 1'),
    (79,  fig2_path,  'Figure 2'),
    (89,  fig3_path,  'Figure 3'),
    (97,  fig4_path,  'Figure 4'),
    (101, fig5_path,  'Figure 5'),
    (109, fig6_path,  'Figure 6'),
    (111, fig7_path,  'Figure 7'),
    (113, fig8_path,  'Figure 8'),
    (122, fig9_path,  'Figure 9'),
    (124, fig10_path, 'Figure 10'),
]

for para_idx, fig_path, fig_name in replacements:
    if not os.path.exists(fig_path):
        print(f"  [SKIP] {fig_name}: file not found ({fig_path})")
        continue
    ok = replace_image_in_docx(DOCX_PATH, para_idx, fig_path)
    print(f"  {'OK  ' if ok else 'FAIL'}: {fig_name} → para {para_idx}")

doc_check = Document(DOCX_PATH)
print(f"\n  inline_shapes in docx: {len(doc_check.inline_shapes)}")
print(f"\n  Docx: {DOCX_PATH}")
print("Done.")
